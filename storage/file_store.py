"""
Versioned file store for tailored resumes and application artifacts.

Directory structure created per application:
    applications/
      {company_slug}/
        {role_slug}/
          {date}_{time}/
            resume.docx           ← tailored Word document
            resume.pdf            ← exported PDF
            job_description.txt   ← raw JD text
            tailoring_metadata.json ← what was extracted + generated

The master resume is NEVER modified. This module always works with copies.
"""

from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional

from loguru import logger

_APPLICATIONS_ROOT = Path("applications")
_MASTER_RESUME_PATH = Path("resumes") / "master_resume.docx"


# ── Slug helpers ──────────────────────────────────────────────────────────────


def _slugify(text: str, max_len: int = 40) -> str:
    """Convert free text to a safe directory name slug."""
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"[\s-]+", "_", text)
    return text[:max_len].strip("_")


def _version_id(company: str, title: str) -> str:
    """Return a human-readable version ID like stripe_backend_engineer_2026-03-08."""
    date_str = datetime.utcnow().strftime("%Y-%m-%d")
    return f"{_slugify(company)}_{_slugify(title)}_{date_str}"


# ── ApplicationDir ─────────────────────────────────────────────────────────────


class ApplicationDir:
    """
    Represents the versioned directory for one job application.

    Create with ApplicationDir.create(company, title) — never instantiate directly.
    """

    def __init__(self, path: Path, version_id: str):
        self.path = path
        self.version_id = version_id

    @classmethod
    def create(cls, company: str, title: str) -> "ApplicationDir":
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H%M%S")
        dir_path = (
            _APPLICATIONS_ROOT
            / _slugify(company)
            / _slugify(title)
            / timestamp
        )
        dir_path.mkdir(parents=True, exist_ok=True)
        version_id = _version_id(company, title)
        logger.debug("Created application dir: {}", dir_path)
        return cls(dir_path, version_id)

    # ── File paths ─────────────────────────────────────────────────────────────

    @property
    def resume_docx(self) -> Path:
        return self.path / "resume.docx"

    @property
    def resume_pdf(self) -> Path:
        return self.path / "resume.pdf"

    @property
    def jd_txt(self) -> Path:
        return self.path / "job_description.txt"

    @property
    def metadata_json(self) -> Path:
        return self.path / "tailoring_metadata.json"

    # ── Write helpers ──────────────────────────────────────────────────────────

    def save_jd(self, description: str) -> None:
        self.jd_txt.write_text(description, encoding="utf-8")

    def save_metadata(self, data: dict) -> None:
        self.metadata_json.write_text(
            json.dumps(data, indent=2, default=str), encoding="utf-8"
        )

    def copy_master_resume(self, master_path: Optional[Path] = None) -> Path:
        """Copy the master resume into this dir as the base for tailoring."""
        src = master_path or _MASTER_RESUME_PATH
        if not src.exists():
            raise FileNotFoundError(
                f"Master resume not found at {src}. "
                "Place your resume at resumes/master_resume.docx"
            )
        shutil.copy2(src, self.resume_docx)
        logger.debug("Copied master resume to {}", self.resume_docx)
        return self.resume_docx


# ── Resume builder ─────────────────────────────────────────────────────────────


class ResumeBuilder:
    """
    Modifies a .docx resume by injecting tailored bullet points
    into the most recent experience section, then exports to PDF.

    The master resume is never touched — all changes go into a copy.

    Usage:
        builder = ResumeBuilder()
        pdf_path = builder.build(app_dir, bullets=["Led...", "Built...", "Reduced..."])
    """

    def build(
        self,
        app_dir: ApplicationDir,
        bullets: list[str],
        master_path: Optional[Path] = None,
    ) -> Path:
        """
        Copy master resume → inject bullets → export PDF.

        Args:
            app_dir:     ApplicationDir instance for this application.
            bullets:     List of 3 tailored bullet point strings.
            master_path: Override the master resume path.

        Returns:
            Path to the generated PDF.
        """
        docx_path = app_dir.copy_master_resume(master_path)
        self._inject_bullets(docx_path, bullets)
        pdf_path = self._export_pdf(docx_path, app_dir.resume_pdf)
        logger.info("Resume built: {}", pdf_path)
        return pdf_path

    def _inject_bullets(self, docx_path: Path, bullets: list[str]) -> None:
        """
        Replace the bullet points in the most recent experience section.

        Strategy: find the first bulleted list in the document (which should be
        the most recent role's responsibilities), replace those paragraphs with
        the new bullets. Preserves all formatting.
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return

        doc = Document(str(docx_path))
        self._replace_first_bullet_block(doc, bullets)
        doc.save(str(docx_path))

    @staticmethod
    def _all_paragraphs(doc):
        """Yield all paragraphs in the document including those inside tables."""
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def _replace_first_bullet_block(self, doc, bullets: list[str]) -> None:
        """
        Find the first block of experience bullet paragraphs in the document
        (including inside tables) and replace their text with the provided bullets.

        Detection order:
          1. Paragraphs with a bullet list style or bullet-prefix character.
          2. Fallback: long Normal-style paragraphs (>40 chars) after a section
             heading that contains "experience".
        """
        bullet_styles = {
            "List Bullet", "List Bullet 2", "List Bullet 3",
            "List Paragraph", "List",
        }

        all_paras = list(self._all_paragraphs(doc))

        # Pass 1 — classic bullet style / bullet prefix character
        block_start: Optional[int] = None
        block_end: Optional[int] = None
        for i, para in enumerate(all_paras):
            is_bullet = (
                para.style.name in bullet_styles
                or para.text.strip().startswith(("•", "·", "-", "*", "○"))
            )
            if is_bullet:
                if block_start is None:
                    block_start = i
                block_end = i
            elif block_start is not None:
                break

        # Pass 2 — fallback for Normal-style table resumes (no explicit bullet
        # markers).  Find "Professional Experience" heading, then collect the
        # immediately following long-text paragraphs as the bullet block.
        if block_start is None:
            experience_idx: Optional[int] = None
            for i, para in enumerate(all_paras):
                if "experience" in para.text.lower() and len(para.text.strip()) < 60:
                    experience_idx = i
                    break

            if experience_idx is not None:
                for i, para in enumerate(all_paras[experience_idx + 1:], start=experience_idx + 1):
                    text = para.text.strip()
                    if len(text) > 40:
                        if block_start is None:
                            block_start = i
                        block_end = i
                    elif block_start is not None:
                        break  # First gap after the bullet block — stop

        if block_start is None:
            logger.warning("No bullet block found in resume — bullets NOT injected")
            return

        # Replace the first N paragraphs with new bullets; leave the rest unchanged.
        # (If the master resume has 8 bullets but we only generated 3, paragraphs
        # 4-8 keep their original text rather than being wiped to empty strings.)
        block_paras = all_paras[block_start:block_end + 1]
        for i, para in enumerate(block_paras):
            if i >= len(bullets):
                break  # leave remaining original bullets as-is
            bullet_text = bullets[i].lstrip("•·-* ").strip()
            if para.runs:
                para.runs[0].text = bullet_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = bullet_text

        logger.debug(
            "Injected {} bullets into resume (block at para {}–{})",
            len(bullets), block_start, block_end,
        )

    def _export_pdf(self, docx_path: Path, pdf_path: Path) -> Path:
        """
        Export .docx to PDF. Tries WeasyPrint → LibreOffice → fallback stub.
        """
        # Try LibreOffice (best quality, available on macOS via brew)
        if self._export_libreoffice(docx_path, pdf_path):
            return pdf_path

        # Try python-docx2pdf
        if self._export_docx2pdf(docx_path, pdf_path):
            return pdf_path

        # Final fallback: use the DOCX directly (Playwright can upload any file type).
        # Log a clear install hint so the user knows how to fix it.
        logger.warning(
            "PDF export failed — falling back to DOCX. "
            "To enable PDF export, install LibreOffice:  brew install --cask libreoffice"
        )
        return docx_path  # caller checks existence; DOCX upload still works

    def _export_libreoffice(self, docx_path: Path, pdf_path: Path) -> bool:
        """Convert via LibreOffice headless (brew install --cask libreoffice)."""
        import subprocess
        import shutil as sh

        soffice = sh.which("soffice") or sh.which("libreoffice")
        if not soffice:
            return False
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(pdf_path.parent), str(docx_path)],
                capture_output=True, timeout=30,
            )
            # LibreOffice outputs to same dir with same stem + .pdf
            generated = pdf_path.parent / (docx_path.stem + ".pdf")
            if generated.exists() and generated != pdf_path:
                generated.rename(pdf_path)
            return pdf_path.exists()
        except Exception as exc:
            logger.debug("LibreOffice export failed: {}", exc)
            return False

    def _export_docx2pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        """Convert via docx2pdf (requires licensed Microsoft Word on macOS)."""
        import io
        import contextlib
        try:
            from docx2pdf import convert  # type: ignore
            # docx2pdf may call sys.exit() on Word licence errors — catch BaseException
            # and suppress its stdout chatter (it prints error dicts directly).
            buf = io.StringIO()
            with contextlib.redirect_stdout(buf):
                convert(str(docx_path), str(pdf_path))
            return pdf_path.exists()
        except SystemExit:
            logger.debug("docx2pdf called sys.exit() — Word likely not licensed")
            return False
        except Exception as exc:
            logger.debug("docx2pdf export failed: {}", exc)
            return False
